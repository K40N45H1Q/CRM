#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from datetime import datetime, timedelta
import os
import subprocess
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FONT_NAME = "FreeSerif"
FONT_SIZE_PT = 9

SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"

class DOCXFiller:
    def __init__(self, template_path: str, output_directory: str = ".", 
                 output_pdf: bool = True, keep_docx: bool = False, **kwargs):
        today = datetime.now().date()
        signing_date = today.strftime("%d.%m.%Y")
        agreement_number = f"AG-{today.year}-{datetime.now().strftime('%H%M%S')}"
        payload = {f"{{{k}}}": str(v) for k, v in kwargs.items()}
        if "{signingDate}" not in payload:
            payload["{signingDate}"] = signing_date
        if "{agreementNumber}" not in payload:
            payload["{agreementNumber}"] = agreement_number
        if "{agreementDuration}" in payload and "{expirationDate}" not in payload:
            try:
                years = int(payload["{agreementDuration}"])
                signing_obj = datetime.strptime(signing_date, "%d.%m.%Y").date()
                try:
                    exp = signing_obj.replace(year=signing_obj.year + years)
                except ValueError:
                    exp = signing_obj + timedelta(days=365 * years)
                payload["{expirationDate}"] = exp.strftime("%d.%m.%Y")
            except Exception:
                payload["{expirationDate}"] = signing_date
        self.payload = payload
        self.document = Document(template_path)
        self._process_paragraphs_and_tables()
        self._process_headers_and_footers()
        self._process_xml_text_nodes()
        docx_path = os.path.join(output_directory, f"{agreement_number}.docx")
        self.pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        self.document.save(docx_path)
        if output_pdf:
            self._convert_to_pdf(docx_path)
            if not keep_docx and os.path.exists(docx_path):
                os.remove(docx_path)

    def _convert_to_pdf(self, docx_path: str):
        if not os.path.exists(SOFFICE):
            print(f"❌ LibreOffice не найден: {SOFFICE}")
            return
        
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        out_dir = os.path.abspath(os.path.dirname(pdf_path) or ".")
        abs_docx = os.path.abspath(docx_path)

        cmd = [
            SOFFICE,
            "--headless",
            "--invisible",
            "--norestore",
            "--nologo",
            "--convert-to", "pdf",
            "--outdir", out_dir,
            abs_docx
        ]

        try:
            proc = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                creationflags=subprocess.CREATE_NO_WINDOW,
                timeout=60
            )
            
            if proc.returncode == 0 and os.path.exists(pdf_path):
                print(f"✅ PDF создан: {pdf_path}")
            else:
                try: err_msg = proc.stderr.decode('cp1251').strip()
                except: err_msg = proc.stderr.decode('utf-8', errors='ignore').strip()
                print(f"❌ Ошибка конвертации: {err_msg if err_msg else 'Код 1'}")

        except subprocess.TimeoutExpired:
            print("❌ Таймаут LibreOffice.")
        except Exception as e:
            print(f"❌ Исключение: {e}")

    def _clear_bold_and_set_fonts_in_rPr(self, rPr_element):
        # Удаляем жирность
        for child in list(rPr_element):
            if child.tag == qn("w:b"): rPr_element.remove(child)
        b = OxmlElement("w:b"); b.set(qn("w:val"), "false"); rPr_element.append(b)
        
        # Устанавливаем шрифт
        rf = OxmlElement("w:rFonts")
        for a in ["w:ascii","w:hAnsi","w:eastAsia","w:cs"]: rf.set(qn(a), FONT_NAME)
        rPr_element.append(rf)
        
        # Устанавливаем размер
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(Pt(FONT_SIZE_PT).pt*2))); rPr_element.append(sz)

    def _apply_format_to_run(self, run, highlight: bool = False):
        """Применяет шрифт и размер. Highlight игнорируется (убран желтый фон)."""
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)
        run.font.bold = False
        rPr = run._element.get_or_add_rPr()
        self._clear_bold_and_set_fonts_in_rPr(rPr)
        # Блок с желтым фоном (shd/highlight) полностью удален

    def _apply_format_to_run_xml(self, run_xml_element, highlight: bool = False):
        """Применяет шрифт и размер к XML элементу. Highlight игнорируется."""
        rPr = run_xml_element.get_or_add_rPr()
        self._clear_bold_and_set_fonts_in_rPr(rPr)
        # Блок с желтым фоном полностью удален

    def _process_paragraphs_and_tables(self):
        paragraphs = list(self.document.paragraphs)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells: paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs: self._replace_in_paragraph_runs(paragraph)

    def _process_headers_and_footers(self):
        for section in self.document.sections:
            for p in section.header.paragraphs: self._replace_in_paragraph_runs(p)
            for t in section.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: self._replace_in_paragraph_runs(p)
            for p in section.footer.paragraphs: self._replace_in_paragraph_runs(p)
            for t in section.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: self._replace_in_paragraph_runs(p)

    def _replace_in_paragraph_runs(self, paragraph):
        if not paragraph.text: return
        runs = paragraph.runs
        run_texts = [r.text or "" for r in runs]
        full_text = "".join(run_texts)
        search_text = " ".join(full_text.split())
        for key, val in self.payload.items():
            key_search = " ".join(key.split())
            pos = 0
            while True:
                idx = search_text.find(key_search, pos)
                if idx == -1: break
                key_nospace = key.replace(" ", "")
                orig_idx = full_text.find(key_nospace) if key_nospace in full_text else full_text.find(key_search)
                if orig_idx == -1: break
                acc = 0; start_run = None
                for i, t in enumerate(run_texts):
                    if acc + len(t) > orig_idx: start_run = i; break
                    acc += len(t)
                end_pos = orig_idx + len(key_nospace)
                acc2 = 0; end_run = None
                for j, t in enumerate(run_texts):
                    if acc2 + len(t) >= end_pos: end_run = j; break
                    acc2 += len(t)
                if start_run is None or end_run is None: break
                
                # Очищаем текст во всех затронутых раннах
                for k in range(start_run, end_run + 1): runs[k].text = ""
                # Вставляем новое значение в первый ранн
                runs[start_run].text = val
                
                # Применяем форматирование (шрифт/размер) ко всем затронутым раннам
                # highlight=False убирает желтый фон
                for k in range(start_run, end_run + 1): 
                    self._apply_format_to_run(runs[k], highlight=False)
                
                pos = idx + len(key_search)
                run_texts = [r.text or "" for r in runs]
                full_text = "".join(run_texts)
                search_text = " ".join(full_text.split())

    def _process_xml_text_nodes(self):
        def replace_in_element(element):
            nodes = list(element.iter(qn("w:t")))
            i = 0
            while i < len(nodes):
                j = i; combined = nodes[j].text or ""; found_key = None
                while True:
                    scombined = " ".join(combined.split())
                    for key in self.payload.keys():
                        if " ".join(key.split()) in scombined: found_key = key; break
                    if found_key or j + 1 >= len(nodes): break
                    j += 1; combined += nodes[j].text or ""
                if not found_key: i = j + 1; continue
                original = "".join([n.text or "" for n in nodes[i:j+1]])
                replaced = " ".join(original.split())
                changed = False
                for key, val in self.payload.items():
                    ksearch = " ".join(key.split())
                    if ksearch in replaced: replaced = replaced.replace(ksearch, val); changed = True
                if changed:
                    nodes[i].text = replaced
                    for extra in nodes[i+1:j+1]: extra.text = ""
                    # highlight=False убирает желтый фон
                    self._apply_format_to_run_xml(nodes[i].getparent(), highlight=False)
                i = j + 1
        replace_in_element(self.document._element)
        for rel in list(self.document.part.rels.values()):
            if "header" in rel.reltype or "footer" in rel.reltype: replace_in_element(rel.target_part.element)